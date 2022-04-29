from multiprocessing import Manager
from django.db import models
from django.forms import ValidationError 
from django.utils.text import slugify
from django.contrib.auth.models import  AbstractUser, User, UserManager as AuthUserManager
import random
from ckeditor.fields import RichTextField
from datetime import datetime, timedelta
from django.db.models import Q
#for documents
from docx import Document
from django.conf import settings

User._meta.get_field('email').blank=True
User._meta.get_field('email')._unique=False
User._meta.get_field('id')._unique=True

class UserManager(AuthUserManager):
    
    def get_active_users(self):
        return self.model.objects.filter(Q(date_joined__gte=datetime.now()-timedelta(days=365))|Q(position__iexact='a')|Q(position__iexact='i'))


class User_data(AbstractUser):
    CHOICE_LIST = [
        ("a", "admin"),
        ("p", "parent"),
        ("s", "student"),
        ("i", "instructor"),
        ("n", "none"),
    ]

    patronymic = models.CharField(max_length=100, null=True, blank=True)
    birthdate = models.DateField(null=True, blank=True)
    phone_number = models.CharField(max_length=100, null=True, blank=True)
    extra_phone_numbers = models.CharField(max_length=100, null=True, blank=True)
    position = models.CharField(max_length=1, choices=CHOICE_LIST, default="n")
    profile_photo = models.ImageField(default="default.jpg", upload_to="user_img",null=True, blank=True)
    notes = models.TextField(null=True, blank=True)
    blocked = models.BooleanField(default=False)
    individual_type = models.BooleanField(default=False, null=True, blank=True)
    passport_address = models.CharField(max_length=300, null=True, blank=True)
    passport_number = models.CharField(max_length=9, null=True, blank=True)
    passport_serial = models.CharField(max_length=4, null=True, blank=True)
    passport_who_give = models.CharField(max_length=100, null=True, blank=True)
    passport_when_give = models.CharField(
        max_length=12, default="0000-00-00", null=True, blank=True
    )
    passport_file = models.FileField(null=True, blank=True, upload_to="passport_copies")
    passport_file1 = models.FileField(
        null=True, blank=True, upload_to="passport_copies"
    )
    office_address = models.CharField(max_length=500, null=True, blank=True)
    office_bank_account = models.CharField(max_length=500, null=True, blank=True)
    office_bank_code = models.CharField(max_length=500, null=True, blank=True)
    office_inn = models.CharField(max_length=500, null=True, blank=True)
    office_licence_file = models.FileField(default=None, upload_to="office_files", blank=True, null=True)
    # create_date = models.DateTimeField(auto_now_add=True)
    update_date = models.DateTimeField(auto_now=True, null=True, blank=True)
    deleted = models.BooleanField(default=False, null=True, blank=True)
    class Meta:
        verbose_name = "User Account"
        verbose_name = "User Accounts"
        ordering = ("-id",)
 
    def __str__(self):
        return f"{self.first_name} {self.last_name} {self.position}"

    @property
    def get_parent(self):
        return self.student.all() or None

    @property
    def get_student(self):
        return self.parent.all() or None

    objects = UserManager()

    @property
    def full_name(self):
        return f"{self.last_name} {self.first_name}"

    def save(self, *args, **kwargs):
        if not self.id:
            self.id = User_data.objects.first().id+1
        if not self.username and self.first_name:
            try:
                self.username = self.first_name
            except Exception as e:
                raise ValidationError('This username has been taken!')
        else:
            ValidationError('This username has been taken!-Это имя пользователя занято! ')
        if not self.password:
            self.set_password(f"{self.first_name}{self.id}")
        if 'pbkdf2_sha256$260000' not in self.password:
            self.set_password(self.password)
        super(User_data, self).save(*args, **kwargs)
    
    # def delete(self, *args, **kwargs):
    #     if self.first_name == "" and self.last_name == "" and self.passport_number == "" and self.passport_serial == "":
    #         return super(User_data, self).delete(*args, **kwargs)



class School_branches(models.Model):
    school_name = models.CharField(max_length=100, null=True, blank=True)
    branch_name = models.CharField(max_length=100, null=True, blank=True)
    address = models.CharField(max_length=1000, null=True, blank=True)
    create_date = models.DateTimeField(auto_now_add=True)
    details = models.TextField(null=True, blank=True)

    class Meta:
        verbose_name = "School branch"
        verbose_name_plural = "School branches"

    def __str__(self):
        return f"{self.branch_name} {self.address}"


class Timeslots(models.Model):
    timeslot_name = models.CharField(max_length=300, null=True, blank=True)
    start_time = models.TimeField()
    end_time = models.TimeField()
    duration = models.IntegerField()
    mon = models.BooleanField(default=False)
    tue = models.BooleanField(default=False)
    wed = models.BooleanField(default=False)
    thu = models.BooleanField(default=False)
    fri = models.BooleanField(default=False)
    sat = models.BooleanField(default=False)
    sun = models.BooleanField(default=False)
    school_name = models.CharField(max_length=100)
    branch_name = models.CharField(max_length=100, null=True, blank=True)

    class Meta:
        verbose_name = "Time Slot"
        verbose_name = "Time Slots"
        ordering = ('-id', )
    def save(self, *args, **kwargs):
        # if not self.id:
        #     self.id = Timeslots.objects.first().id+1
        return super(Timeslots, self).save(*args, **kwargs)

    def __str__(self):
        return self.timeslot_name


class Call_orders(models.Model):
    phone_number = models.CharField(max_length=45)
    order_time = models.DateTimeField(auto_now_add=True, null=True, blank=True)
    confirmed = models.BooleanField(default=False)
    confirm_time = models.DateTimeField(null=True, blank=True)

    class Meta:
        verbose_name = "Call order"
        verbose_name_plural = "Call orders"

    def __str__(self):
        return self.phone_number


class Classrooms(models.Model):
    building = models.CharField(max_length=10)
    room = models.CharField(max_length=20, null=True, blank=True)
    capacity = models.IntegerField(null=True, blank=True)
    details = models.CharField(max_length=300, null=True, blank=True)
    school_name = models.CharField(max_length=100, default="Codecraft.uz")
    branch_name = models.CharField(max_length=100)
    create_date = models.DateTimeField(auto_now_add=True)

    class Meta:
        verbose_name = "Classroom"
        verbose_name_plural = "Classrooms"

    def __str__(self):
        return f"{self.building} {self.room} {self.capacity}"


class Company_Data(models.Model):
    data_field = models.CharField(max_length=100)
    value = models.CharField(max_length=1000)
    last_change = models.DateTimeField(auto_now=True, null=True, blank=True)

    class Meta:
        verbose_name = "Company Data"
        verbose_name_plural = "Company Datas"

    def __str__(self):
        return self.data_field


#
# 
#
class Courses(models.Model):
    short_title = models.CharField(null=True, blank=True, max_length=100)
    title = models.CharField(null=True, blank=True, max_length=300)
    price = models.IntegerField(null=True, blank=True)
    subtitle = models.CharField(null=True, blank=True, max_length=300)
    lessons = models.IntegerField()
    lesson_duration = models.IntegerField()
    min_students = models.IntegerField(null=True, blank=True)
    max_students = models.IntegerField(null=True, blank=True)
    description_file = models.FileField(null=True, blank=True, upload_to="Course_files")
    title_image = models.FileField(null=True, blank=True, upload_to="Course_files")
    cover_image = models.FileField(null=True, blank=True, upload_to="Course_files")
    deleted = models.IntegerField(default=0)
    notes = RichTextField(config_name='default', null=True, blank=True)
    course_section_id = models.IntegerField(null=True, blank=True)
    publicized = models.BooleanField(default=False)
    curriculum = RichTextField(config_name='default', null=True, blank=True)  #models.TextField(null=True, blank=True)
    required_course_id = models.IntegerField(default=1, null=True, blank=True)
    video_link = models.URLField(null=True, blank=True)
    slug = models.SlugField(max_length=100, null=True, blank=True)

    class Meta:
        verbose_name = "Course"
        verbose_name_plural = "Courses"

    def __str__(self):
        return self.short_title

    def save(self, *args, **kwargs):
        if not self.slug:
            self.slug = slugify(self.short_title)
        # if not self.id:
        #     self.id = Courses.objects.last().id+1
        super(Courses, self).save(*args, **kwargs)
    
    @property
    def get_groups(self):
        groups = self.groups.all()
        return groups.filter(finished__isnull=True)


class Courses_Slugs(models.Model):
    slug = models.CharField(max_length=100)
    slug_name = models.CharField(max_length=100)
    themify_icon = models.CharField(max_length=100)

    def __str__(self):
        return self.slug


class Courses_Slugs_id(models.Model):
    course = models.ForeignKey(Courses, on_delete=models.CASCADE)
    slug = models.ForeignKey(Courses_Slugs, on_delete=models.CASCADE)

    def __str__(self):
        return f"{self.course.short_title} {self.slug.slug_name}"


class Course_Sections(models.Model):
    short_title = models.CharField(max_length=100, null=True)
    title = models.CharField(max_length=500, null=True, blank=True)
    short_description = models.TextField(null=True, blank=True)
    head_instructor = models.ForeignKey(
        User_data, on_delete=models.CASCADE, null=True, blank=True
    )
    description = models.TextField(max_length=400, null=True, blank=True)
    image_file = models.FileField(default="default.jpg", upload_to="courses_sections")

    def __str__(self):
        return f"{self.short_title}"


class GroupManager(models.Manager):

    def get_active_groups(self):
        return super().get_queryset().filter(finished__isnull=True)
    def get_archive_groups(self):
        return super().get_queryset().filter(finished__isnull=False)

class Group(models.Model):
    name = models.CharField(max_length=200, null=True, blank=True)
    course = models.ForeignKey(Courses, on_delete=models.CASCADE, related_name='groups')
    timeslot = models.ForeignKey(Timeslots, on_delete=models.CASCADE)
    start_date = models.DateField()
    end_date = models.DateField()
    classroom_building = models.CharField(max_length=10)
    classroom_room = models.CharField(max_length=8)
    finished = models.DateTimeField(null=True, blank=True)
    finish = models.BooleanField(default=False)
    activate_again = models.BooleanField(default=False)
    def __str__(self):
        return f"{self.name}"

    class Meta:
        ordering = ("-id",)

    @property
    def get_instructor(self):
        # if self.group_instructor.first().instructor:
        return [obj.instructor for obj in self.group_instructor.all()]
        # return ''
    @property
    def get_group_students(self):
        return self.group_students.all()


    objects = GroupManager()

    def save(self, *args, **kwargs):
        # if not self.id:
        #     self.id = Group.objects.first().id+1
        if self.finish:
            self.finished = datetime.now()
        if self.finish and self.activate_again:
            self.finished = None
        return super(Group, self).save(*args, **kwargs)


class GroupInstructors(models.Model):
    group = models.ForeignKey(Group, on_delete=models.CASCADE, related_name="group_instructor")
    instructor = models.ForeignKey(User_data, on_delete=models.SET_NULL, null=True)

    def __str__(self):
        return f"{self.group.id} {self.instructor.first_name}"


def re_generate_contract_file(contract_data=None):
    if contract_data is not None:
        document = Document(settings.MEDIA_ROOT+"/"+"contract_template.docx")
        for p in document.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                for key in contract_data.keys():
                    if key in text:
                        text=text.replace(key,contract_data[key])
                        inline[i].text = text
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key in contract_data.keys():
                            if key in cell.text:
                                inline = p.runs
                                for i in range(len(inline)):
                                    if key in inline[i].text:
                                        text = inline[i].text.replace(key, contract_data[key])
                                        inline[i].text = text
        doc_name = settings.MEDIA_ROOT+'/contracts/'+f"{contract_data['contract_serial']}{contract_data['last_contract_no']}.docx"
        import os
        if os.path.exists(doc_name):   #==> remove old contract
            os.remove(doc_name)
        
        document.save(doc_name)
        doc_name_media = f"/contracts/{contract_data['contract_serial']}{contract_data['last_contract_no']}.docx"
        return doc_name_media
    return None

def generate_contract_file(contract_data=None):
    if contract_data is not None:
        document = Document(settings.MEDIA_ROOT+"/"+"contract_template.docx")
        for p in document.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                for key in contract_data.keys():
                    if key in text:
                        text=text.replace(key,contract_data[key])
                        inline[i].text = text
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key in contract_data.keys():
                            if key in cell.text:
                                inline = p.runs
                                for i in range(len(inline)):
                                    if key in inline[i].text:
                                        text = inline[i].text.replace(key, contract_data[key])
                                        inline[i].text = text
        doc_name = settings.MEDIA_ROOT+'/contracts/'+f"{contract_data['contract_serial']}{contract_data['last_contract_no']}.docx"
        document.save(doc_name)
        
        doc_name_media = f"/contracts/{contract_data['contract_serial']}{contract_data['last_contract_no']}.docx"
        
        return doc_name_media
    else:
        return None


class Group_Students(models.Model):
    student = models.ForeignKey(
        User_data, on_delete=models.SET_NULL, null=True, blank=True, related_name="student_groups"
    )
    group = models.ForeignKey(Group, on_delete=models.CASCADE, related_name="group_students")
    points = models.IntegerField(null=True, blank=True)
    certificate = models.FileField(upload_to="group_students", null=True, blank=True)
    contract_no = models.CharField(max_length=20, null=True, blank=True)
    contract = models.FileField(upload_to="contracts", null=True, blank=True)
    discount = models.IntegerField(null=True, blank=True)
    confirmed = models.BooleanField(default=False, help_text='Check this for creating new contract.')
    re_gen_contract = models.BooleanField(default=False, help_text='Check this for regenerating new contract.')
    class Meta:
        ordering = ('-id',)
    def __str__(self):
        return f"{self.student.first_name}{self.student.last_name} {self.group.id}"

    def save(self, *args, **kwargs):
        company_data = Company_Data.objects.all()
        contract_data = {}
        for comp_data in company_data:
            contract_data[comp_data.data_field] = comp_data.value

        if self.student.position == "s":
            try:
                parent = Parenting_students.objects.get(student=self.student).parent
                if parent is None:
                    parent = self.student
            except:
                parent = self.student
            
            if parent.full_name:
                contract_data["parent_full_name"] = parent.full_name
            else:
                contract_data["parent_full_name"] = " "
            if parent.passport_serial:
                contract_data["passport_serial"] = parent.passport_serial
            else:
                contract_data["passport_serial"] = " "
            if parent.passport_number:
                contract_data["passport_number"] = parent.passport_number
            else:
                contract_data["passport_number"] = " "
            if parent.passport_who_give:
                contract_data["passport_who_give"] = parent.passport_who_give
            else:
                contract_data["passport_who_give"] = " "
            if parent.passport_when_give:
                contract_data["passport_when_give"] = parent.passport_when_give
            else:
                contract_data["passport_when_give"] = " "
            if parent.passport_address:
                contract_data["passport_address"] = parent.passport_address
            else:
                contract_data["passport_address"] = " "
            if parent.phone_number:
                contract_data["parent_phone_number"] = parent.phone_number
            else:
                contract_data["parent_phone_number"] = " "
            if self.student.full_name:
                contract_data["student_full_name"] = self.student.full_name
            else:
                contract_data["student_full_name"] = " "
            if self.student.individual_type:
                contract_data["price"] = str(self.group.course.price)
            else:
                contract_data["price"] = str(self.group.course.price)
            if self.student.birthdate:
                contract_data["b_year"] = f"{self.student.birthdate.year}"

            now = datetime.now()

            contract_data["day"] = "0"+str(now.day) if now.day<10 else str(now.day)
            contract_data["month"] = "0"+str(now.month) if now.month<10 else str(now.month)
            contract_data["cur_year"] = str(now.year)

            contract_data["hours"] = str(int(self.group.course.lessons)*2)
            contract_data["short_title"] = self.group.course.short_title
            
            #regenerate contract ----------------
            if (self.contract_no or self.contract) and self.re_gen_contract:
                try:
                    contract_data["last_contract_no"] = str(self.contract_no[2:])
                except:
                    import os
                    if os.exists(self.contract):
                        contract_data["last_contract_no"] = "".join(os.path.split(self.contract)[1].split('.')[0])
                    else:
                        raise FileNotFoundError('File not found!')
                self.contract = re_generate_contract_file(contract_data=contract_data)
                self.contract_no = contract_data["contract_serial"]+contract_data["last_contract_no"]
            
            # generate contract ------------------
            elif self.confirmed and (not self.contract) and (not self.contract_no):
                print("Contract data: ", contract_data)
                contract_data["last_contract_no"] = str(int(contract_data["last_contract_no"])+1)
                self.contract = generate_contract_file(contract_data=contract_data)
                self.contract_no = contract_data["contract_serial"]+contract_data["last_contract_no"]
                try:
                    last_contract_no = Company_Data.objects.get(data_field='last_contract_no')
                    last_contract_no.value=str(contract_data['last_contract_no'])
                    last_contract_no.save()
                except Exception as e:
                    raise e
            # -------------------------------
                
            # if not self.id:
            #     self.id = Group_Students.objects.last().id
            return super(Group_Students, self).save(*args, **kwargs)

class Parenting_students(models.Model):
    student = models.ForeignKey(User_data, on_delete=models.SET_NULL, related_name="student", null=True)
    parent = models.ForeignKey(User_data, on_delete=models.SET_NULL, related_name="parent", null=True)

    def __str__(self):
        if self.student:
            return (
                f"{self.student.full_name}"
            )
        return f"{self.id} None"


PERCENTAGES = (
    ('5', "5 % chegirma"),
    ('10', "10 % chegirma"),
    ('15', "15 % chegirma"),
    ('20', "20 % chegirma"),
    ('25', "25 % chegirma"),
    ('30', "30 % chegirma"),
    ('35', "35 % chegirma"),
    ('40', "40 % chegirma"),
    ('50', "50 % chegirma"),
)
MONTH = (
    (1, "January"),
    (2, "February"),
    (3, "March"),
    (4, "April"),
    (5, "May"),
    (6, "June"),
    (7, "July"),
    (8, "August"),
    (9, "September"),
    (10, "October"),
    (11, "Nowember"),
    (12, "December"),
)
class Payments(models.Model):
    parent = models.ForeignKey(
        User_data,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name="parent_payment",
    )
    student = models.ForeignKey(
        User_data,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name="student_payment",
    )
    group = models.ForeignKey(Group, on_delete=models.CASCADE)
    payment_method = models.CharField(max_length=30, default="CASH")
    amount = models.IntegerField()
    date_contracted = models.DateTimeField(auto_now_add=True)
    date_time = models.DateTimeField(auto_now=True)  # auto_now_add=True
    deleted = models.BooleanField(default=False)
    for_month = models.IntegerField(null=True, blank=True, choices=MONTH)
    admin = models.ForeignKey(
        User_data,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name="admin_account",
    )
    notes = models.CharField(max_length=200, default="--blocked")
    paid = models.BooleanField(default=False)

    def __str__(self):
        return f"{self.id} {self.date_time}"


    class Meta:
        ordering = ('-date_time',)

class Prereq(models.Model):
    course = models.ForeignKey(
        Courses, on_delete=models.CASCADE, related_name="main_course"
    )
    required_course = models.ForeignKey(
        Courses, on_delete=models.CASCADE, related_name="required_course"
    )

    def __str__(self):
        return str(self.id)


class WhyUs(models.Model):
    title = models.CharField(max_length=1000)
    image = models.FileField(upload_to="FrontendFiles", null=True, blank=True)
    icon = models.CharField(max_length=300, null=True, blank=True)
    description = RichTextField(config_name='default', null=True, blank=True)

    def __str__(self):
        return f"{self.title}"


class FAQ(models.Model):
    title = models.CharField(max_length=1000)
    description = RichTextField(config_name='default', null=True, blank=True)

    def __str__(self):
        return f"{self.title}"


class FrontPages(models.Model):
    title = models.CharField(max_length=1000)
    description = RichTextField(config_name='default', null=True, blank=True)
    image = models.ImageField(upload_to="FrontendFiles")

    def __str__(self):
        return f"{self.title}"


class FrontStatistics(models.Model):
    title = models.CharField(max_length=1000)
    value = models.CharField(max_length=100)

    def __str__(self):
        return f"{self.title} {self.value}"


def get_random():
    code = ""
    for i in range(6):
        code += str(random.randint(0, 9))
    return code


class VerificationCode(models.Model):
    email = models.EmailField()
    code = models.CharField(max_length=6, default=get_random())
    status = models.BooleanField(default=False)

    class Meta:
        verbose_name = "VerificationCode"

    def __str__(self):
        return f"{self.email} : {self.code}"


class Prices(models.Model):
    course_type = models.CharField(max_length=200, null=True, blank=True)
    USD_online = models.CharField(max_length=10, null=True, blank=True)
    UZS_online = models.CharField(max_length=10, null=True, blank=True)
    EUR_online = models.CharField(max_length=10, null=True, blank=True)
    RUBL_online = models.CharField(max_length=10, null=True, blank=True)
    USD_individual = models.CharField(max_length=10, null=True, blank=True)
    UZS_individual = models.CharField(max_length=10, null=True, blank=True)
    EUR_individual = models.CharField(max_length=10, null=True, blank=True)
    RUBL_individual = models.CharField(max_length=10, null=True, blank=True)
    USD_offline = models.CharField(max_length=10, null=True, blank=True)
    UZS_offline = models.CharField(max_length=10, null=True, blank=True)
    EUR_offline = models.CharField(max_length=10, null=True, blank=True)
    RUBL_offline = models.CharField(max_length=10, null=True, blank=True)

    def __str__(self):
        return f"{self.course_type}"


class Actions(models.Model):
    title = models.CharField(max_length=400, null=True, blank=True)
    text = RichTextField(config_name='default', null=True, blank=True)
    image = models.ImageField(upload_to="action_images")
    date_added = models.DateField(auto_now_add=True)

    def __str__(self):
        return f"{self.title}"


class Contact(models.Model):
    phone = models.CharField(max_length=20, null=True, blank=True)
    instagram = models.URLField(null=True, blank=True)
    facebook = models.URLField(null=True, blank=True)
    telegram = models.URLField(null=True, blank=True)

    def __str__(self) -> str:
        return f"Company Social Info"


class SigningUpFreeLesson(models.Model):
    first_name = models.CharField(max_length=100)
    last_name = models.CharField(max_length=120)
    phone = models.CharField(max_length=50)
    sent_to_tg_bot = models.BooleanField(default=False)

    def __str__(self):
        return f"{self.first_name} {self.last_name}"


#Barno opam uchun:
class StudentGroupModels(models.Model):
    student_id=models.IntegerField()
    group = models.ForeignKey(Group, on_delete=models.CASCADE)



class CompanyArchives(models.Model):
    name = models.CharField(max_length=500, help_text="Tarix uchun eslatma qoldiring.", null=True, blank=True)
    file = models.FileField(upload_to="archives", null=True, blank=True)
    date_start = models.DateField(help_text="Bu yerda ma\'lumotlarni arxivga saqlash uchun boshlang\'ich sana")
    date_end = models.DateField(help_text="Bu yerda ma\'lumotlarni arxivga saqlash uchun yakuniy sana")
    admin = models.CharField(max_length=500, help_text="Bu yerda arxivga kiritgan adminning ismi.", null=True, blank=True)
    date_archived = models.DateField(auto_now_add=True)

    class Meta:
        verbose_name = "Company Archieve Info"
        verbose_name_plural = "Company Archive Infos"
    
    def __str__(self):
        return f"{self.name} archived at {self.date_archived} by {self.admin}."



REBATE = (
    (5, "5 % chegirma"),
    (10, "10 % chegirma"),
    (15, "15 % chegirma"),
    (20, "20 % chegirma"),
    (25, "25 % chegirma"),
    (30, "30 % chegirma"),
    (35, "35 % chegirma"),
    (40, "40 % chegirma"),
    (45, "45 % chegirma"),
    (50, "50 % chegirma"),
)

class Rebate(models.Model):
    student   = models.ForeignKey(User_data, on_delete=models.SET_NULL, null=True, blank=True)
    group     = models.ForeignKey(Group, on_delete=models.SET_NULL, null=True, blank=True)
    notes     = models.CharField(max_length=100, null=True, blank=True)
    rebate    = models.IntegerField(default=0, choices=REBATE)
    for_month = models.IntegerField(null=True, blank=True, choices=MONTH)
    date_added = models.DateTimeField(auto_now_add=True)

    class Meta:
        verbose_name        = "Chegirma"
        verbose_name_plural = "Chegirmalar"
        ordering = ("-date_added",)

    def __str__(self):
        return self.student.full_name 
